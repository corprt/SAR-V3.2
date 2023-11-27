import random,os,json,io,re,zipfile,tempfile
import numpy as np
import pandas as pd
import ssl
import streamlit as st
from io import BytesIO
from io import StringIO
import streamlit.components.v1 as components
from typing import List, Dict, Any
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import streamlit as st
import streamlit_toggle as tog
from langchain import HuggingFaceHub
from langchain.llms import OpenAI
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
import openai 
import fitz
import docx
from gtts import gTTS
import PyPDF2
from PyPDF2 import PdfReader
from langchain import PromptTemplate, LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.chains import ConversationChain
from langchain.memory import ConversationBufferMemory
from langchain.memory import ConversationSummaryBufferMemory
from langchain.chains.conversation.prompt import ENTITY_MEMORY_CONVERSATION_TEMPLATE
from langchain.chains.conversation.memory import ConversationEntityMemory
from langchain.callbacks import get_openai_callback
from langchain.chains import RetrievalQA
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from usellm import Message, Options, UseLLM
from huggingface_hub import login
from io import StringIO
from io import BytesIO
import cv2
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from fpdf import FPDF
import base64
from langchain.document_transformers import (
    LongContextReorder,
)

@st.cache_data(show_spinner=False)
def usellm(prompt):
    """
    Getting GPT-3.5 Model into action
    """
    service = UseLLM(service_url="https://usellm.org/api/llm")
    messages = [
      Message(role="system", content="You are a fraud analyst, who is an expert at finding out suspicious activities"),
      Message(role="user", content=f"{prompt}"),
      ]
    options = Options(messages=messages)
    response = service.chat(options)
    return response.content

def print_docs(docs):
    st.write(f"\n{'-' * 100}\n".join([f"Document {i+1}:\n\n" + d.page_content for i, d in enumerate(docs)]))

@st.cache_data
def show_pdf(file_path):
    with open(file_path,"rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
        pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="1000px" type="application/pdf"></iframe>'
        st.markdown(pdf_display, unsafe_allow_html=True)

@st.cache_data
def pdf_to_bytes(pdf_file_):
    with open(pdf_file_,"rb") as pdf_file:
        pdf_content = pdf_file.read()
        pdf_bytes_io = io.BytesIO(pdf_content)
    return pdf_bytes_io

@st.cache_data
def read_pdf_files(path):
    pdf_files =[]
    directoty_path = path
    files = os.listdir(directoty_path)
    for file in files:
            pdf_files.append(file)
    return pdf_files


@st.cache_data
def merge_pdfs(pdf_list):
    """
    Helper function to merge PDFs
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        pdf_document = PyPDF2.PdfReader(pdf)
        pdf_merger.append(pdf_document)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    return output_pdf


@st.cache_data
def process_text(text):
    # Add your custom text processing logic here
    processed_text = text
    return processed_text



    
@st.cache_data
def merge_and_extract_text(pdf_list):
    """
    Helper function to merge PDFs and extract text
    """
    pdf_merger = PyPDF2.PdfMerger()
    for pdf in pdf_list:
        with open(pdf, 'rb') as file:
            pdf_merger.append(file)
    output_pdf = BytesIO()
    pdf_merger.write(output_pdf)
    pdf_merger.close()
    
    # Extract text from merged PDF
    merged_pdf = PyPDF2.PdfReader(output_pdf)
    all_text = []
    for page in merged_pdf.pages:
        text = page.extract_text()
        all_text.append(text)
    
    return ' '.join(all_text)

def reset_session_state():
    session_state = st.session_state
    session_state.clear()


@st.cache_data
def render_pdf_as_images(pdf_file):
    """
    Helper function to render PDF pages as images
    """
    pdf_images = []
    pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        img = page.get_pixmap()
        img_bytes = img.tobytes()
        pdf_images.append(img_bytes)
    pdf_document.close()
    return pdf_images

# To check if pdf is searchable
def is_searchable_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True

    return False


@st.cache_data(show_spinner=False)
def extract_text_from_pdf(file_path):
    with pdfplumber.open(file_path) as pdf:
        all_text = []
        for page in pdf.pages:
            text = page.extract_text()
            all_text.append(text)
    return "\n".join(all_text)

@st.cache_data
def add_footer_with_fixed_text(doc, footer_text):
    # Create a footer object
    footer = doc.sections[0].footer

    # Add a paragraph to the footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    # Set the fixed text in the footer
    paragraph.text = footer_text

    # Add a page number field to the footer
    run = paragraph.add_run()
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
    fld_simple = parse_xml(fld_xml)
    run._r.append(fld_simple)
    # Set the alignment of the footer text
    paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    
@st.cache_data
def create_filled_box_with_text(color, text):
    box_html = f'<div style="flex: 1; height: 100px; background-color: {color}; display: flex; align-items: center; justify-content: center;">{text}</div>'
    st.markdown(box_html, unsafe_allow_html=True)


# Function to add checkboxes to the DataFrame
@st.cache_data
def add_checkboxes_to_dataframe(df):
    # Create a new column 'Select' with checkboxes
    checkbox_values = [True] * (len(df) - 1) + [False]  # All True except the last row
    df['Select'] = checkbox_values
    return df

# convert scanned pdf to searchable pdf
# @st.cache_data(show_spinner=False)
def convert_scanned_pdf_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    images = convert_from_path(input_file)

    # Preprocess images using OpenCV
    for i, image in enumerate(images):
        # Convert image to grayscale
        image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

        # Apply thresholding to remove noise
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

        # Enhance contrast
        image = cv2.equalizeHist(image)

        # Save preprocessed image
        cv2.imwrite(f'{i}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    for i in range(len(images)):
        image = cv2.imread(f'{i}.png')
        text += pytesseract.image_to_string(image)
    
    return text




def st_audiorec():

    # get parent directory relative to current directory
    parent_dir = os.path.dirname(os.path.abspath(__file__))
    # Custom REACT-based component for recording client audio in browser
    build_dir = os.path.join(parent_dir, "st_audiorec/frontend/build")
    # specify directory and initialize st_audiorec object functionality
    st_audiorec = components.declare_component("st_audiorec", path=build_dir)

    # Create an instance of the component: STREAMLIT AUDIO RECORDER
    raw_audio_data = st_audiorec()  # raw_audio_data: stores all the data returned from the streamlit frontend
    wav_bytes = None                # wav_bytes: contains the recorded audio in .WAV format after conversion

    # the frontend returns raw audio data in the form of arraybuffer
    # (this arraybuffer is derived from web-media API WAV-blob data)

    if isinstance(raw_audio_data, dict):  # retrieve audio data
        with st.spinner('retrieving audio-recording...'):
            ind, raw_audio_data = zip(*raw_audio_data['arr'].items())
            ind = np.array(ind, dtype=int)  # convert to np array
            raw_audio_data = np.array(raw_audio_data)  # convert to np array
            sorted_ints = raw_audio_data[ind]
            stream = BytesIO(b"".join([int(v).to_bytes(1, "big") for v in sorted_ints]))
            # wav_bytes contains audio data in byte format, ready to be processed further
            wav_bytes = stream.read()

    return wav_bytes


# @st.cache_data
def text_to_docs(text: str,filename) -> List[Document]:
    """Converts a string or list of strings to a list of Documents
    with metadata."""
    
    if isinstance(text, str):
        # Take a single string as one page
        text = [text]
    page_docs = [Document(page_content=page) for page in text]

    # Add page numbers as metadata
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    # Split pages into chunks
    doc_chunks = []

    for doc in page_docs:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=450,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=50,
        )
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, 
                metadata = {
                "page": i + 1,"chunk": i} )
            # Add sources a metadata
            doc.metadata["source"] = filename
            doc_chunks.append(doc)
    return doc_chunks

# def convert_scanned_pdf_to_searchable_pdf(input_file, output_file):
#     """
#      Convert a Scanned PDF to Searchable PDF

#     """
#     # Convert PDF to images
#     images = convert_from_path(input_file)

#     # Preprocess images using OpenCV
#     for i, image in enumerate(images):
#         # Convert image to grayscale
#         image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

#         # Apply thresholding to remove noise
#         _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

#         # Enhance contrast
#         image = cv2.equalizeHist(image)

#         # Save preprocessed image
#         cv2.imwrite(f'{i}.png', image)

#     # Perform OCR on preprocessed images using Tesseract
#     text = ''
#     for i in range(len(images)):
#         image = cv2.imread(f'{i}.png')
#         text += pytesseract.image_to_string(image)

#     # Add searchable layer to PDF using PyPDF2
#     pdf_writer = PyPDF2.PdfFileWriter()
#     with open(input_file, 'rb') as f:
#         pdf_reader = PyPDF2.PdfFileReader(f)
#         for i in range(pdf_reader.getNumPages()):
#             page = pdf_reader.getPage(i)
#             pdf_writer.addPage(page)
#             pdf_writer.addBookmark(f'Page {i+1}', i)

#     pdf_writer.addMetadata({
#         '/Title': os.path.splitext(os.path.basename(input_file))[0],
#         '/Author': 'Doc Manager',
#         '/Subject': 'Searchable PDF',
#         '/Keywords': 'PDF, searchable, OCR',
#         '/Creator': 'Py script',
#         '/Producer': 'EXL Service',
#     })

#     pdf_writer.addAttachment('text.txt', text.encode())

#     with open(output_file, 'wb') as f:
#         pdf_writer.write(f)

#     # Clean up temporary files
#     for i in range(len(images)):
#         os.remove(f'{i}.png')

@st.cache_data(show_spinner=False)
def convert_image_to_searchable_pdf(input_file):
    """
     Convert a Scanned PDF to Searchable PDF

    """
    # Convert PDF to images
    # images = convert_from_path(input_file)

    # # Preprocess images using OpenCV
    # for i, image in enumerate(input_file):
    # Convert image to grayscale
    image = cv2.imread(input_file)
    image = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)

    # Apply thresholding to remove noise
    _, image = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Enhance contrast
    image = cv2.equalizeHist(image)

    
    file = os.path.basename(input_file)
    # Save preprocessed image
    cv2.imwrite(f'{input_file}.png', image)

    # Perform OCR on preprocessed images using Tesseract
    text = ''
    # for i in range(len(input_file)):
    image = cv2.imread(f'{input_file}.png')
    text += pytesseract.image_to_string(image)

    return text


@st.cache_data(show_spinner=False)
def LLM_Response(query,context,prompt,llm):
    llm_chain = LLMChain(prompt,llm)
    response = llm_chain.run({"query":query, "context":context})
    return response

llama_13b = HuggingFaceHub(
            repo_id="meta-llama/Llama-2-13b-chat-hf",
            model_kwargs={"temperature":0.01, 
                        "min_new_tokens":100, 
                        "max_new_tokens":300})

@st.cache_data(show_spinner=False)
def llama_llm(_llm,prompt):
    response = _llm.predict(prompt)
    return response

# Chunking with overlap
text_splitter = RecursiveCharacterTextSplitter(
chunk_size = 1000,
chunk_overlap  = 100,
length_function = len,
separators=["\n\n", "\n", " ", ""]
)



# @st.cache_data(show_spinner=False)
def embed(model_name):
    hf_embeddings = HuggingFaceEmbeddings(model_name=model_name)
    return hf_embeddings


#@st.cache_data
# def embedding_store(temp_file_path,hf_embeddings):
#     merged_pdf = merge_pdfs(temp_file_path)
#     final_pdf = PyPDF2.PdfReader(merged_pdf)
#     text = ""
#     for page in final_pdf.pages:
#         text += page.extract_text()

#     texts =  text_splitter.split_text(text)
#     docs = text_to_docs(texts)
#     docsearch = FAISS.from_documents(docs, hf_embeddings)
#     return docs, docsearch


# @st.cache_data(show_spinner=False)
# def embedding_store(text,_hf_embeddings):
#     texts =  text_splitter.split_text(text)
#     docs = []
#     for i, chunk in enumerate(texts):
#         doc = Document(
#             page_content = chunk,
#             metadata = {
#                 "page": i + 1,
#                 "chunk": i}
#         )
#         # Add sources a metadata
#         doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
#         docs.append(doc)
#     docs = text_to_docs(texts)
#     docsearch = FAISS.from_documents(docs, _hf_embeddings)
#     return docs,docsearch

    
# @st.cache_data(show_spinner=False)
def embedding_store(_doc,_hf_embeddings):
    docsearch = FAISS.from_documents(_doc, _hf_embeddings)
    return _doc, docsearch










